"""
Word 简历模板填充脚本
用法: python fill_docx.py <template> <resume_json> <output_path>

template: 模板名称 (cn-classic-blue / cn-sidebar-navy)
resume_json: 简历数据 JSON 文件路径
output_path: 输出 .docx 文件路径

JSON 数据格式见 fill_docx_schema.json
"""

import json
import sys
import os
from pathlib import Path
from copy import deepcopy
from docx import Document
from lxml import etree

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
WPS_NS = 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape'

TEMPLATE_DIR = Path(__file__).parent.parent.parent.parent / 'templates' / 'word'


def get_textboxes(body):
    return list(body.iter(f'{{{WPS_NS}}}txbx'))


def set_textbox_text(txbx, lines, preserve_first_run_format=True):
    w_p = f'{{{W_NS}}}p'
    w_r = f'{{{W_NS}}}r'
    w_t = f'{{{W_NS}}}t'
    w_rPr = f'{{{W_NS}}}rPr'
    w_pPr = f'{{{W_NS}}}pPr'
    w_br = f'{{{W_NS}}}br'

    txbx_content = txbx[0]

    first_para = txbx_content.find(w_p)
    if first_para is None:
        return

    first_run = first_para.find(w_r)
    run_format = None
    if first_run is not None:
        rPr = first_run.find(w_rPr)
        if rPr is not None:
            run_format = deepcopy(rPr)

    para_format = None
    pPr = first_para.find(w_pPr)
    if pPr is not None:
        para_format = deepcopy(pPr)

    for p in list(txbx_content.findall(w_p)):
        txbx_content.remove(p)

    for i, line in enumerate(lines):
        new_p = etree.SubElement(txbx_content, w_p)
        if para_format is not None and i == 0:
            new_p.insert(0, deepcopy(para_format))

        new_r = etree.SubElement(new_p, w_r)
        if run_format is not None:
            new_r.insert(0, deepcopy(run_format))

        new_t = etree.SubElement(new_r, w_t)
        new_t.text = line
        new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')


def fill_cn_classic_blue(doc, data):
    body = doc.element.body
    tbs = get_textboxes(body)

    name = data.get('name', '')
    age = data.get('age', '')
    edu_level = data.get('education_level', '')
    job_target = data.get('job_target', '')

    set_textbox_text(tbs[11], [
        f'姓    名：  {name}',
        f'年    龄：  {age}',
        f'学    历：  {edu_level}',
        f'求职意向：  {job_target}',
    ])

    phone = data.get('phone', '')
    email = data.get('email', '')
    wechat = data.get('wechat', '')
    address = data.get('address', '')
    set_textbox_text(tbs[9], [
        f'手机：  {phone}',
        f'邮箱：  {email}',
        f'微信：  {wechat}',
        f'地址：  {address}',
    ])

    self_eval = data.get('self_evaluation', '')
    set_textbox_text(tbs[8], [self_eval])

    edu_lines = []
    for edu in data.get('education', []):
        header = f"{edu.get('major', '')}    {edu.get('school', '')}（{edu.get('degree', '')}）    {edu.get('dates', '')}"
        edu_lines.append(header)
        if edu.get('courses'):
            edu_lines.append(f"主修课程：{edu['courses']}")
    set_textbox_text(tbs[6], edu_lines)

    work = data.get('work_experience', [])
    if len(work) >= 1:
        w = work[0]
        lines = [f"{w.get('title', '')}    {w.get('company', '')}    {w.get('dates', '')}"]
        for d in w.get('duties', []):
            lines.append(d)
        set_textbox_text(tbs[4], lines)

    if len(work) >= 2:
        w = work[1]
        lines = [f"{w.get('title', '')}    {w.get('company', '')}    {w.get('dates', '')}"]
        for d in w.get('duties', []):
            lines.append(d)
        set_textbox_text(tbs[2], lines)

    skills = data.get('skills', [])
    if skills:
        set_textbox_text(tbs[1], skills)


def fill_cn_sidebar_navy(doc, data):
    body = doc.element.body
    tbs = get_textboxes(body)

    set_textbox_text(tbs[6], [data.get('name', '')])

    age = data.get('age', '')
    address = data.get('address', '')
    phone = data.get('phone', '')
    email = data.get('email', '')
    set_textbox_text(tbs[7], [age, address, phone, email])

    set_textbox_text(tbs[12], [f"求职意向：{data.get('job_target', '')}"])

    top_skills = data.get('top_skills', data.get('skills', []))[:4]
    skill_boxes = [tbs[8], tbs[9], tbs[10], tbs[11]]
    for i, sb in enumerate(skill_boxes):
        if i < len(top_skills):
            set_textbox_text(sb, [top_skills[i]])

    edu_lines = []
    for edu in data.get('education', []):
        header = f"{edu.get('dates', '')}    {edu.get('school', '')}    {edu.get('major', '')}（{edu.get('degree', '')}）"
        edu_lines.append(header)
        if edu.get('courses'):
            edu_lines.append(edu['courses'])
    set_textbox_text(tbs[13], edu_lines)

    work_lines = []
    for w in data.get('work_experience', []):
        work_lines.append(f"{w.get('dates', '')}    {w.get('company', '')}    {w.get('title', '')}")
        for d in w.get('duties', []):
            work_lines.append(d)
    set_textbox_text(tbs[15], work_lines)

    awards = data.get('awards', [])
    if awards:
        set_textbox_text(tbs[17], awards)

    self_eval = data.get('self_evaluation', '')
    if self_eval:
        set_textbox_text(tbs[19], [self_eval])


def replace_cell_text(cell, new_text):
    lines = new_text.split('\n') if new_text else ['']
    paras = cell.paragraphs

    for i, line in enumerate(lines):
        if i < len(paras):
            _set_para_text(paras[i], line)
        else:
            _add_para_to_cell(cell, line, paras[-1] if paras else None)

    tc = cell._element
    w_p = f'{{{W_NS}}}p'
    all_p = tc.findall(w_p)
    for p in all_p[len(lines):]:
        tc.remove(p)


def _set_para_text(para, text):
    for run in para.runs:
        run.text = ''
    if para.runs:
        para.runs[0].text = text
    else:
        run = para.add_run(text)


def _add_para_to_cell(cell, text, template_para):
    from copy import deepcopy
    w_p = f'{{{W_NS}}}p'
    w_pPr = f'{{{W_NS}}}pPr'
    w_r = f'{{{W_NS}}}r'
    w_rPr = f'{{{W_NS}}}rPr'
    w_t = f'{{{W_NS}}}t'

    new_p = etree.SubElement(cell._element, w_p)
    if template_para is not None:
        pPr = template_para._element.find(w_pPr)
        if pPr is not None:
            new_p.insert(0, deepcopy(pPr))
        first_run = template_para._element.find(w_r)
        rPr = None
        if first_run is not None:
            rPr = first_run.find(w_rPr)

    new_r = etree.SubElement(new_p, w_r)
    if template_para is not None and first_run is not None and rPr is not None:
        new_r.insert(0, deepcopy(rPr))
    new_t = etree.SubElement(new_r, w_t)
    new_t.text = text
    new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')


def _remove_para(para):
    p = para._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def replace_para_text(para, new_text):
    _set_para_text(para, new_text)


def fill_cn_translator(doc, data):
    paras = doc.paragraphs

    replace_para_text(paras[0], data.get('name', ''))

    phone = data.get('phone', '')
    email = data.get('email', '')
    replace_para_text(paras[1], f'手机：{phone} | 邮箱：{email}')

    domains = data.get('translation_domains', '英汉双语翻译')
    replace_para_text(paras[2], domains)

    job_target = data.get('job_target', '')
    replace_para_text(paras[3], f'求职意向：{job_target}')

    tables = doc.tables

    self_eval = data.get('self_evaluation', '')
    if self_eval:
        replace_cell_text(tables[0].rows[0].cells[0], self_eval)

    edu = data.get('education', [])
    if len(edu) >= 1:
        e = edu[0]
        replace_cell_text(tables[1].rows[1].cells[0],
            f"{e.get('major', '')} ｜ {e.get('degree', '')}\n{e.get('school', '')}")
        replace_cell_text(tables[1].rows[1].cells[1], e.get('dates', ''))
        details = f"GPA：{e.get('gpa', '')}"
        if e.get('honors'):
            details += f"\n荣誉：{e['honors']}"
        replace_cell_text(tables[1].rows[2].cells[0], details)

    if len(edu) >= 2:
        e = edu[1]
        txt = f"{e.get('major', '')} ｜ {e.get('degree', '')}\n{e.get('school', '')}"
        if e.get('gpa'):
            txt += f"\nGPA：{e['gpa']}"
        if e.get('honors'):
            txt += f"\n荣誉：{e['honors']}"
        replace_cell_text(tables[1].rows[3].cells[0], txt)
        replace_cell_text(tables[1].rows[3].cells[1], e.get('dates', ''))

    work = data.get('work_experience', [])
    work_table = tables[2]
    for i, w in enumerate(work):
        row_header = i * 2
        row_detail = i * 2 + 1
        if row_header >= len(work_table.rows):
            break
        replace_cell_text(work_table.rows[row_header].cells[0],
            f"{w.get('title', '')}\n{w.get('company', '')}")
        replace_cell_text(work_table.rows[row_header].cells[1],
            f"{w.get('dates', '')}\n{w.get('city', '')}")
        if row_detail < len(work_table.rows):
            duties = '\n'.join(w.get('duties', []))
            replace_cell_text(work_table.rows[row_detail].cells[0], duties)

    activities = data.get('activities', [])
    act_table = tables[3]
    if len(activities) >= 1:
        a = activities[0]
        replace_cell_text(act_table.rows[0].cells[0],
            f"{a.get('title', '')}\n{a.get('organization', '')}")
        replace_cell_text(act_table.rows[0].cells[1], a.get('dates', ''))
        if 1 < len(act_table.rows):
            duties = '\n'.join(a.get('duties', []))
            replace_cell_text(act_table.rows[1].cells[0], duties)
    if len(activities) >= 2:
        a = activities[1]
        content = f"{a.get('title', '')}\n{a.get('organization', '')}"
        duties = a.get('duties', [])
        if duties:
            content += '\n' + '\n'.join(duties)
        replace_cell_text(act_table.rows[2].cells[0], content)
        replace_cell_text(act_table.rows[2].cells[2], a.get('dates', ''))

    certs = data.get('certificates', '')
    tools = data.get('tools', '')
    cert_text = f"证书：{certs}\n软件工具：{tools}"
    replace_cell_text(tables[4].rows[0].cells[0], cert_text)

    _auto_fit_row_heights(doc)


def _auto_fit_row_heights(doc):
    w_trPr = f'{{{W_NS}}}trPr'
    w_trHeight = f'{{{W_NS}}}trHeight'
    for table in doc.tables:
        for row in table.rows:
            trPr = row._tr.find(w_trPr)
            if trPr is not None:
                trHeight = trPr.find(w_trHeight)
                if trHeight is not None:
                    trPr.remove(trHeight)


def fill_en_traditional(doc, data):
    paras = doc.paragraphs

    replace_para_text(paras[0], data.get('name', ''))

    contact_parts = []
    if data.get('email'):
        contact_parts.append(data['email'])
    if data.get('phone'):
        contact_parts.append(data['phone'])
    replace_para_text(paras[1], ' ● '.join(contact_parts))

    replace_para_text(paras[2], data.get('address', ''))

    edu = data.get('education', [])
    if edu:
        e = edu[0]
        replace_para_text(paras[6], e.get('school', ''))
        degree_line = e.get('degree', '')
        if e.get('dates'):
            degree_line += ', ' + e['dates']
        replace_para_text(paras[7], degree_line)
        replace_para_text(paras[8], 'Relevant courses: ' + e.get('courses', '') if e.get('courses') else '')

    paras_to_remove = []

    work = data.get('work_experience', [])
    if len(work) >= 1:
        w = work[0]
        header = f"{w.get('company', '')} | {w.get('dates', '')}"
        replace_para_text(paras[12], header)
        replace_para_text(paras[13], w.get('title', ''))
        duties = w.get('duties', [])
        duty_start = 16
        for di, d in enumerate(duties):
            idx = duty_start + di
            if idx < len(paras):
                replace_para_text(paras[idx], d)
        for idx in range(duty_start + len(duties), 22):
            if idx < len(paras):
                paras_to_remove.append(paras[idx])

    if len(work) >= 2:
        w = work[1]
        header = f"{w.get('company', '')} | {w.get('dates', '')}"
        replace_para_text(paras[23], header)
        replace_para_text(paras[24], w.get('title', ''))
        duties = w.get('duties', [])
        duty_start = 27
        for di, d in enumerate(duties):
            idx = duty_start + di
            if idx < len(paras):
                replace_para_text(paras[idx], d)
        for idx in range(duty_start + len(duties), 33):
            if idx < len(paras):
                paras_to_remove.append(paras[idx])

    skills = data.get('skills', [])
    if skills and len(paras) > 37:
        replace_para_text(paras[37], skills[0] if len(skills) > 0 else '')
    if len(skills) > 1 and len(paras) > 38:
        replace_para_text(paras[38], skills[1] if len(skills) > 1 else '')
    interests = data.get('interests', '')
    if len(paras) > 39:
        replace_para_text(paras[39], 'Interest: ' + interests if interests else '')

    for p in paras_to_remove:
        _remove_para(p)


def fill_en_contempo(doc, data):
    paras = doc.paragraphs

    replace_para_text(paras[0], data.get('name', ''))

    contact = data.get('address', '')
    if data.get('phone'):
        contact += ' | ' + data['phone']
    if data.get('email'):
        contact += ' | ' + data['email']
    replace_para_text(paras[1], contact)

    replace_para_text(paras[3], data.get('summary', ''))

    tables = doc.tables
    skills = data.get('skills', [])
    mid = len(skills) // 2
    if tables:
        replace_cell_text(tables[0].rows[0].cells[0], '\n'.join(skills[:mid]) if skills else '')
        replace_cell_text(tables[0].rows[0].cells[1], '\n'.join(skills[mid:]) if skills else '')

    paras_to_remove = []

    edu = data.get('education', [])
    if edu:
        e = edu[0]
        edu_text = e.get('degree', '')
        if e.get('dates'):
            edu_text += ', ' + e['dates']
        replace_para_text(paras[6], edu_text)
        replace_para_text(paras[7], e.get('school', ''))
        if len(paras) > 8:
            replace_para_text(paras[8], e.get('certifications', '') if e.get('certifications') else '')
        for idx in range(9, 11):
            if idx < len(paras):
                paras_to_remove.append(paras[idx])

    work = data.get('work_experience', [])
    if len(work) >= 1:
        w = work[0]
        replace_para_text(paras[12], f"{w.get('title', '')}, {w.get('dates', '')}")
        replace_para_text(paras[13], f"{w.get('company', '')}")
        duties = w.get('duties', [])
        for di, d in enumerate(duties):
            idx = 14 + di
            if idx < 19:
                replace_para_text(paras[idx], d)
        for idx in range(14 + len(duties), 19):
            if idx < len(paras):
                paras_to_remove.append(paras[idx])

    if len(work) >= 2:
        w = work[1]
        replace_para_text(paras[19], f"{w.get('title', '')}, {w.get('dates', '')}")
        replace_para_text(paras[20], f"{w.get('company', '')}")
        duties = w.get('duties', [])
        for di, d in enumerate(duties):
            idx = 21 + di
            if idx < 24:
                replace_para_text(paras[idx], d)
        for idx in range(21 + len(duties), 24):
            if idx < len(paras):
                paras_to_remove.append(paras[idx])

    langs = data.get('languages', [])
    lang_start = 69
    for li, lang in enumerate(langs):
        idx = lang_start + li
        if idx < len(paras):
            replace_para_text(paras[idx], lang)

    last_work_idx = 23 if len(work) >= 2 else 18
    for idx in range(last_work_idx + 1, lang_start):
        if idx < len(paras):
            paras_to_remove.append(paras[idx])
    lang_end = lang_start + len(langs)
    for idx in range(lang_end, len(paras)):
        paras_to_remove.append(paras[idx])

    for p in paras_to_remove:
        _remove_para(p)


def fill_en_executive(doc, data):
    paras = doc.paragraphs
    tables = doc.tables

    replace_para_text(paras[0], data.get('name', ''))
    replace_para_text(paras[1], data.get('address', ''))

    contact = ''
    if data.get('phone'):
        contact += data['phone']
    if data.get('email'):
        contact += ' - ' + data['email']
    replace_para_text(paras[2], contact)

    replace_para_text(paras[4], data.get('summary', ''))

    skills = data.get('skills', [])
    mid = len(skills) // 2
    if tables:
        replace_cell_text(tables[0].rows[0].cells[0], '\n'.join(skills[:mid]) if skills else '')
        replace_cell_text(tables[0].rows[0].cells[1], '\n'.join(skills[mid:]) if skills else '')

    edu = data.get('education', [])
    if edu and len(tables) > 1:
        e = edu[0]
        edu_text = e.get('degree', '')
        if e.get('school'):
            edu_text += '\n' + e['school']
        if e.get('certifications'):
            edu_text += '\n' + e['certifications']
        replace_cell_text(tables[1].rows[0].cells[1], edu_text)

    work = data.get('work_experience', [])
    if len(work) >= 1 and len(tables) > 2:
        w = work[0]
        replace_cell_text(tables[2].rows[0].cells[0], w.get('dates', ''))
        content = w.get('title', '') + '\n' + w.get('company', '')
        for d in w.get('duties', []):
            content += '\n' + d
        replace_cell_text(tables[2].rows[0].cells[1], content)

    if len(work) >= 2 and len(tables) > 3:
        w = work[1]
        replace_cell_text(tables[3].rows[0].cells[0], w.get('dates', ''))
        content = w.get('title', '') + '\n' + w.get('company', '')
        for d in w.get('duties', []):
            content += '\n' + d
        replace_cell_text(tables[3].rows[0].cells[1], content)

    keep_indices = {0, 1, 2, 3, 4, 5, 6, 7}
    for idx in range(len(paras) - 1, -1, -1):
        if idx not in keep_indices:
            _remove_para(paras[idx])

    _auto_fit_row_heights(doc)


def _replace_merged_row(table, row_idx, col_date, col_content, date_text, content_text):
    row = table.rows[row_idx]
    seen = set()
    for ci, cell in enumerate(row.cells):
        cell_id = id(cell._element)
        if cell_id in seen:
            continue
        seen.add(cell_id)
        if ci <= col_date:
            replace_cell_text(cell, date_text)
        else:
            replace_cell_text(cell, content_text)


def fill_en_harvard(doc, data):
    t0 = doc.tables[0]

    replace_cell_text(t0.rows[0].cells[0], data.get('name', ''))
    replace_cell_text(t0.rows[1].cells[0], data.get('title', ''))

    row3 = t0.rows[3]
    seen3 = set()
    parts = {'phone': data.get('phone', ''), 'address': data.get('address', ''), 'email': data.get('email', '')}
    field_order = ['phone', 'address', 'email']
    fi = 0
    for ci, cell in enumerate(row3.cells):
        cid = id(cell._element)
        if cid in seen3:
            continue
        seen3.add(cid)
        if cell.text.strip() or fi < len(field_order):
            if fi < len(field_order):
                replace_cell_text(cell, parts[field_order[fi]])
                fi += 1

    edu = data.get('education', [])
    edu_rows = [7, 9, 11]
    for i, e in enumerate(edu):
        if i >= len(edu_rows):
            break
        ri = edu_rows[i]
        if ri < len(t0.rows):
            _replace_merged_row(t0, ri, 1, 3, e.get('dates', ''),
                e.get('degree', '') + '\n' + e.get('school', '') +
                ('\n' + e.get('details', '') if e.get('details') else ''))

    work = data.get('work_experience', [])
    work_rows = [16, 18]
    for i, w in enumerate(work):
        if i >= len(work_rows):
            break
        ri = work_rows[i]
        if ri < len(t0.rows):
            content = w.get('title', '') + '\n' + w.get('company', '')
            for d in w.get('duties', []):
                content += '\n' + d
            _replace_merged_row(t0, ri, 1, 3, w.get('dates', ''), content)

    if len(doc.tables) > 1:
        t1 = doc.tables[1]

        awards = data.get('awards', [])
        award_rows = [21, 23]
        for i, a in enumerate(awards):
            if i >= len(award_rows):
                break
            ri = award_rows[i]
            if ri < len(t1.rows):
                _replace_merged_row(t1, ri, 0, 2, a.get('date', ''), a.get('description', ''))

        skills_text = '\n'.join(data.get('skills', []))
        if 28 < len(t1.rows):
            replace_cell_text(t1.rows[28].cells[0], skills_text)

    _auto_fit_row_heights(doc)


def fill_en_standout(doc, data):
    tables = doc.tables

    if tables:
        header_cell = tables[0].rows[0].cells[1]
        paras = header_cell.paragraphs
        if paras:
            replace_para_text(paras[0], data.get('name', ''))
        contact = data.get('address', '')
        if data.get('phone'):
            contact += '  ' + data['phone']
        if data.get('email'):
            contact += '  ' + data['email']
        if len(paras) > 1:
            replace_para_text(paras[1], contact)

    if len(tables) > 1:
        summary_cell = tables[1].rows[0].cells[1]
        nested = summary_cell.tables
        if nested:
            target = nested[0].rows[0].cells[1] if len(nested[0].rows[0].cells) > 1 else nested[0].rows[0].cells[0]
            replace_cell_text(target, data.get('summary', ''))

    if len(tables) > 2:
        qual_cell = tables[2].rows[0].cells[1]
        nested = qual_cell.tables
        skills = data.get('skills', [])
        mid = len(skills) // 2
        if nested:
            inner = nested[0]
            if inner.rows[0].cells[0].tables:
                double_nested = inner.rows[0].cells[0].tables[0]
                replace_cell_text(double_nested.rows[0].cells[0], '\n'.join(skills[:mid]))
                replace_cell_text(double_nested.rows[0].cells[1], '\n'.join(skills[mid:]))
            else:
                replace_cell_text(inner.rows[0].cells[0], '\n'.join(skills[:mid]))
                if len(inner.rows[0].cells) > 1:
                    replace_cell_text(inner.rows[0].cells[1], '\n'.join(skills[mid:]))

    if len(tables) > 3:
        edu_cell = tables[3].rows[0].cells[1]
        nested = edu_cell.tables
        edu = data.get('education', [])
        if nested and edu:
            e = edu[0]
            target = nested[0].rows[0].cells[1] if len(nested[0].rows[0].cells) > 1 else nested[0].rows[0].cells[0]
            edu_text = e.get('degree', '') + '\n' + e.get('school', '')
            if e.get('certifications'):
                edu_text += '\n' + e['certifications']
            replace_cell_text(target, edu_text)

    if len(tables) > 4:
        work_cell = tables[4].rows[0].cells[1]
        nested_tables = work_cell.tables
        work = data.get('work_experience', [])
        for i, w in enumerate(work):
            if i >= len(nested_tables):
                break
            nt = nested_tables[i]
            target = nt.rows[0].cells[1] if len(nt.rows[0].cells) > 1 else nt.rows[0].cells[0]
            content = f"{w.get('title', '')}  {w.get('dates', '')}"
            content += f"\n{w.get('company', '')}"
            for d in w.get('duties', []):
                content += '\n' + d
            replace_cell_text(target, content)

    _auto_fit_row_heights(doc)


def fill_en_refined_black(doc, data):
    paras = doc.paragraphs
    tables = doc.tables

    replace_para_text(paras[0], data.get('name', ''))

    contact_parts = []
    if data.get('address'):
        contact_parts.append(data['address'])
    if data.get('phone'):
        contact_parts.append(data['phone'])
    if data.get('email'):
        contact_parts.append(data['email'])
    replace_para_text(paras[4], ' ♦ '.join(contact_parts))

    replace_para_text(paras[7], data.get('summary', ''))

    skills = data.get('skills', [])
    mid = len(skills) // 2
    if tables:
        replace_cell_text(tables[0].rows[0].cells[0], '\n'.join(skills[:mid]) if skills else '')
        replace_cell_text(tables[0].rows[0].cells[1], '\n'.join(skills[mid:]) if skills else '')

    edu = data.get('education', [])
    if edu:
        e = edu[0]
        replace_para_text(paras[10], e.get('degree', ''))
        replace_para_text(paras[11], e.get('school', ''))
        replace_para_text(paras[12], e.get('certifications', '') if e.get('certifications') else '')

    paras_to_remove = []
    for idx in range(13, 15):
        if idx < len(paras):
            paras_to_remove.append(paras[idx])

    work = data.get('work_experience', [])
    if len(work) >= 1:
        w = work[0]
        replace_para_text(paras[16], f"{w.get('title', '')}, {w.get('dates', '')}")
        replace_para_text(paras[17], f"{w.get('company', '')}")
        duties = w.get('duties', [])
        for di, d in enumerate(duties):
            idx = 18 + di
            if idx < 23:
                replace_para_text(paras[idx], d)
        for idx in range(18 + len(duties), 23):
            if idx < len(paras):
                paras_to_remove.append(paras[idx])

    if len(work) >= 2:
        w = work[1]
        replace_para_text(paras[23], f"{w.get('title', '')}, {w.get('dates', '')}")
        replace_para_text(paras[24], f"{w.get('company', '')}")
        duties = w.get('duties', [])
        for di, d in enumerate(duties):
            idx = 25 + di
            if idx < 28:
                replace_para_text(paras[idx], d)
        for idx in range(25 + len(duties), 28):
            if idx < len(paras):
                paras_to_remove.append(paras[idx])

    for idx in range(28, 69):
        if idx < len(paras):
            paras_to_remove.append(paras[idx])

    langs = data.get('languages', [])
    if 70 < len(paras):
        paras_to_remove.append(paras[70])
    for li, lang in enumerate(langs):
        idx = 71 + li
        if idx < len(paras):
            replace_para_text(paras[idx], lang)
    for idx in range(71 + len(langs), 75):
        if idx < len(paras):
            paras_to_remove.append(paras[idx])

    awards = data.get('awards', [])
    for ai, a in enumerate(awards):
        idx = 76 + ai
        if idx < len(paras):
            replace_para_text(paras[idx], a.get('description', ''))
    for idx in range(76 + len(awards), 78):
        if idx < len(paras):
            paras_to_remove.append(paras[idx])

    for idx in range(78, len(paras)):
        paras_to_remove.append(paras[idx])

    for p in paras_to_remove:
        _remove_para(p)


TEMPLATES = {
    'cn-classic-blue': {
        'file': 'cn-classic-blue.docx',
        'fill': fill_cn_classic_blue,
        'name_cn': '经典蓝灰职业风',
    },
    'cn-sidebar-navy': {
        'file': 'cn-sidebar-navy.docx',
        'fill': fill_cn_sidebar_navy,
        'name_cn': '深蓝左右分栏风',
    },
    'cn-translator': {
        'file': 'cn-translator.docx',
        'fill': fill_cn_translator,
        'name_cn': '中文翻译简约风',
    },
    'en-traditional': {
        'file': 'en-traditional-cv.docx',
        'fill': fill_en_traditional,
        'name_cn': 'Traditional 简约英文',
    },
    'en-contempo': {
        'file': 'en-contempo_success_steel_CV.docx',
        'fill': fill_en_contempo,
        'name_cn': 'Contempo 钢灰英文',
    },
    'en-executive': {
        'file': 'en-executive_ecru_CV.docx',
        'fill': fill_en_executive,
        'name_cn': 'Executive 米色英文',
    },
    # TODO: 排版问题待修复
    # 'en-harvard': {
    #     'file': 'en-Harvard-深蓝色.docx',
    #     'fill': fill_en_harvard,
    #     'name_cn': 'Harvard 深蓝学术风',
    # },
    # 'en-standout': {
    #     'file': 'en-standout_CV.docx',
    #     'fill': fill_en_standout,
    #     'name_cn': 'Standout 左栏标题风',
    # },
    'en-refined-black': {
        'file': 'en-refined_black_CV.docx',
        'fill': fill_en_refined_black,
        'name_cn': 'Refined Black 黑色精致风',
    },
}


def fill_template(template_name, data, output_path):
    if template_name not in TEMPLATES:
        print(f"Error: Unknown template '{template_name}'")
        print(f"Available: {', '.join(TEMPLATES.keys())}")
        sys.exit(1)

    tmpl = TEMPLATES[template_name]
    template_path = TEMPLATE_DIR / tmpl['file']

    if not template_path.exists():
        print(f"Error: Template file not found: {template_path}")
        sys.exit(1)

    doc = Document(str(template_path))
    tmpl['fill'](doc, data)

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    print(f"✅ Word 简历已生成: {output}")
    print(f"   模板: {tmpl['name_cn']} ({template_name})")


def main():
    if len(sys.argv) < 4:
        print(__doc__)
        print("\nAvailable templates:")
        for k, v in TEMPLATES.items():
            print(f"  {k} — {v['name_cn']}")
        sys.exit(1)

    template_name = sys.argv[1]
    json_path = sys.argv[2]
    output_path = sys.argv[3]

    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    fill_template(template_name, data, output_path)


if __name__ == '__main__':
    main()
